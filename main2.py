from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, CallbackQueryHandler, PicklePersistence, MessageHandler, Filters, ConversationHandler
from docx import Document
from android import TOKEN
import os

ENTER_NAME, ENTER_PHONE, UPLOAD_RESUME, GET_DOCUMENT = range(4)

def start(update, context):
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text="Привет! Пожалуйста, введите ваше имя:")
    context.user_data[chat_id] = {}
    context.user_data[chat_id]['files'] = []
    return ENTER_NAME

def receive_name(update, context):
    chat_id = update.effective_chat.id
    name = update.message.text
    context.user_data[chat_id]['name'] = name
    context.bot.send_message(chat_id=chat_id, text="Отлично! Теперь введите ваш номер телефона:")
    return ENTER_PHONE

def receive_phone(update, context):
    chat_id = update.effective_chat.id
    phone = update.message.text
    context.user_data[chat_id]['phone'] = phone

    context.bot.send_message(chat_id=chat_id, text="Спасибо! Теперь я отправлю вам несколько файлов:")

    
    for file in os.listdir():
        if file.startswith('admin_'):
            with open(file, 'rb') as f:
                context.bot.send_document(chat_id=chat_id, document=f)

    context.bot.send_message(chat_id=chat_id, text="Пожалуйста, отправьте ваше резюме в расширении.docx:")
    return UPLOAD_RESUME

def receive_document(update, context):
    chat_id = update.effective_chat.id
    document = update.message.document
    file_id = document.file_id
    new_file = context.bot.get_file(file_id)
    file_path = f'{chat_id}_file_{len(context.user_data[chat_id]["files"])}.docx'
    new_file.download(file_path)
    context.user_data[chat_id]['files'].append(file_path)

    
    if len(context.user_data[chat_id]['files']) == 1:
        context.bot.send_message(chat_id=chat_id, text="Файл успешно получен! Отправьте еще файл или /submit для завершения.")
    
    return UPLOAD_RESUME


def create(update, context):
    chat_id = update.effective_chat.id
    if str(chat_id) != '-945828460': 
        return
    context.bot.send_message(chat_id=chat_id, text="Пожалуйста, отправьте .docx файлы:")
    context.user_data[chat_id] = {}
    context.user_data[chat_id]['files'] = []
    return GET_DOCUMENT

def receive_create_document(update, context):
    chat_id = update.effective_chat.id
    document = update.message.document
    file_id = document.file_id
    new_file = context.bot.get_file(file_id)
    file_path = f'admin_{len(os.listdir())}_file.docx' 
    new_file.download(file_path)
    context.bot.send_message(chat_id=chat_id, text="Файл успешно получен! Отправьте еще файл или /submit для завершения.")
    return GET_DOCUMENT

def submit(update, context):
    chat_id = update.effective_chat.id
    send_data_to_me(context, chat_id)
    return ConversationHandler.END

def merge_documents(filepaths):
    merged_document = Document()

    for index, filepath in enumerate(filepaths):
        sub_document = Document(filepath)

    
        merged_document.add_heading(f'Глава {index+1}: {filepath}', level=1)

        
        for element in sub_document.element.body:
            merged_document.element.body.append(element)

    return merged_document

def send_data_to_me(context, chat_id):
    user_data = context.user_data.get(chat_id, None)
    if user_data:
        name = user_data['name']
        phone = user_data['phone']
        files = user_data['files']
        score = len(files)  

        if score >= 5:  
            context.bot.send_message(chat_id=-945828460, text=f"Имя:{name}\nТелефон: {phone}\nБаллы: {score}")

            
            merged_document = merge_documents(files)
            merged_filename = f'{name}_merged_resume.docx'
            merged_document.save(merged_filename)

            with open(merged_filename, 'rb') as f:
                context.bot.send_document(chat_id=-945828460, document=f)

updater = Updater(TOKEN, persistence=PicklePersistence(filename='bot_data'))
dispatcher = updater.dispatcher

conv_handler = ConversationHandler(
    entry_points=[CommandHandler('start', start), CommandHandler('create', create)],
    states={
        ENTER_NAME: [MessageHandler(Filters.text, receive_name)],
        ENTER_PHONE: [MessageHandler(Filters.text, receive_phone)],
        UPLOAD_RESUME: [MessageHandler(Filters.document, receive_document)],
        GET_DOCUMENT: [MessageHandler(Filters.document, receive_create_document)],
    },
    fallbacks=[CommandHandler('start', start), CommandHandler('submit', submit)],
)

dispatcher.add_handler(conv_handler)
dispatcher.add_handler(CommandHandler('create', create))

conv_handler = ConversationHandler(
    entry_points=[CommandHandler('start', start)],
    states={
        ENTER_NAME: [MessageHandler(Filters.text, receive_name)],
        ENTER_PHONE: [MessageHandler(Filters.text, receive_phone)],
        UPLOAD_RESUME: [MessageHandler(Filters.document, receive_document)],
        GET_DOCUMENT: [MessageHandler(Filters.document, receive_create_document)],
    },
    fallbacks=[CommandHandler('start', start), CommandHandler('submit', submit)],
)
#kkkkkkkkkkkkkkkkk
dispatcher.add_handler(conv_handler)

updater.start_polling()
updater.idle() 
