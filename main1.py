from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, CallbackQueryHandler, PicklePersistence, MessageHandler, Filters, ConversationHandler
from android import TOKEN
from docx import Document
import os

ENTER_NAME, ENTER_PHONE, UPLOAD_RESUME, GET_DOCUMENT = range(4)

def start(update, context):
    chat_id = update.effective_chat.id
    context.bot.send_message(chat_id=chat_id, text="Привет! Пожалуйста, введите ваше имя:")
    context.user_data[chat_id] = {}
    context.user_data[chat_id]['files'] = []
    return ENTER_NAME

def receive_name(update, context):
    chat_id = update.effective_chat.id
    name = update.message.text
    context.user_data[chat_id]['name'] = name
    context.bot.send_message(chat_id=chat_id, text="Отлично! Теперь введите ваш номер телефона:")
    return ENTER_PHONE

def receive_phone(update, context):
    chat_id = update.effective_chat.id
    phone = update.message.text
    context.user_data[chat_id]['phone'] = phone
    context.bot.send_message(chat_id=chat_id, text="Спасибо! Пожалуйста, отправьте ваше резюме в расширении.docx:")
    return UPLOAD_RESUME

def receive_document(update, context):
    chat_id = update.effective_chat.id
    document = update.message.document
    file_id = document.file_id
    new_file = context.bot.get_file(file_id)
    file_path = f'{chat_id}_file_{len(context.user_data[chat_id]["files"])}.docx'
    new_file.download(file_path)
    context.user_data[chat_id]['files'].append(file_path)
    context.bot.send_message(chat_id=chat_id, text="Файл успешно получен! Отправьте еще файл или /submit для завершения.")
    return UPLOAD_RESUME

def create(update, context):
    chat_id = update.effective_chat.id
    if str(chat_id) != '-945828460': 
        return
    context.bot.send_message(chat_id=chat_id, text="Пожалуйста, отправьте .docx файлы:")
    context.user_data[chat_id] = {}
    context.user_data[chat_id]['files'] = []
    return GET_DOCUMENT

def receive_create_document(update, context):
    chat_id = update.effective_chat.id
    document = update.message.document
    file_id = document.file_id
    new_file = context.bot.get_file(file_id)
    file_path = f'admin_file_{len(context.user_data[chat_id]["files"])}.docx'
    new_file.download(file_path)
    context.user_data[chat_id]['files'].append(file_path)
    context.bot.send_message(chat_id=chat_id, text="Файл успешно получен!")

def clear_files(update, context):
    chat_id = update.effective_chat.id
    if str(chat_id) != '-945828460':  
        return
    for file in os.listdir():
        if file.startswith('admin_') or file.endswith('.docx'):  
            os.remove(file)
    context.bot.send_message(chat_id=chat_id, text="Все файлы были успешно удалены.")

def submit(update, context):
    chat_id = update.effective_chat.id
    send_data_to_me(context, chat_id)
    return ConversationHandler.END

def send_data_to_me(context, chat_id):
    user_data = context.user_data.get(chat_id, None)
    if user_data:
        name = user_data['name']
        phone = user_data['phone']
        files = user_data['files']
        score = len(files)

        if score >= 5:  
            doc = Document()
            for file in files:
                doc.add_heading(file, level=1)
                subdoc = Document(file)
                for element in subdoc.element.body:
                    doc.element.body.append(element)
            doc.save(f'{name}_compiled.docx')
            context.bot.send_message(chat_id=-945828460, text=f"Имя:{name}\nТелефон: {phone}\nБаллы: {score}")
            with open(f'{name}_compiled.docx', 'rb') as f:
                context.bot.send_document(chat_id=-945828460, document=f)

updater = Updater(TOKEN, persistence=PicklePersistence(filename='bot_data'))
dispatcher = updater.dispatcher

dispatcher.add_handler(CommandHandler('create', create))
dispatcher.add_handler(CommandHandler('clear', clear_files))

conv_handler = ConversationHandler(
    entry_points=[CommandHandler('start', start)],
    states={
        ENTER_NAME: [MessageHandler(Filters.text, receive_name)],
        ENTER_PHONE: [MessageHandler(Filters.text, receive_phone)],
        UPLOAD_RESUME: [MessageHandler(Filters.document, receive_document)],
        GET_DOCUMENT: [MessageHandler(Filters.document, receive_create_document)],
    },
    fallbacks=[CommandHandler('start', start), CommandHandler('submit', submit)],
)

dispatcher.add_handler(conv_handler)

updater.start_polling()
updater.idle()